"""
Genere les 10 fichiers .docx "solution" des exercices supplementaires.
Utilise python-docx pour une construction fidele (vrai fichier Word valide).
Certains elements avances (table des matieres, formule de somme) sont inseres
comme de vrais champs Word (fldChar/instrText) : ils s'affichent correctement
une fois le document ouvert et les champs actualises dans Word reel (etape
suivante du pipeline, via automatisation COM PowerShell).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "solutions_docx")
ASSETS = os.path.join(HERE, "assets")
os.makedirs(OUT, exist_ok=True)


# ---------------------------------------------------------------------------
# Utilitaires bas niveau
# ---------------------------------------------------------------------------

def add_field(paragraph, field_code, bold=False, size=None, cached_text=None):
    """Insere un vrai champ Word (ex: TOC, =SUM(ABOVE)) dans un paragraphe.

    `cached_text` (optionnel) insere une valeur mise en cache entre les
    marqueurs separate/end, exactement comme Word le fait lui-meme apres
    un calcul : le champ reste vivant et recalculable dans Word reel, tout
    en affichant un resultat correct immediatement, y compris dans un
    rendu HTML (pandoc) qui n'execute pas la logique des champs.
    """
    def new_run():
        r = paragraph.add_run()
        if bold:
            r.bold = True
        if size:
            r.font.size = Pt(size)
        return r

    r_begin = new_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r_begin._r.append(fldChar_begin)

    r_instr = new_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    r_instr._r.append(instrText)

    r_sep = new_run()
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    r_sep._r.append(fldChar_sep)

    if cached_text:
        cache_run = new_run()
        cache_run.text = cached_text

    r_end = new_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r_end._r.append(fldChar_end)

    return r_begin


def set_cell_shading(cell, hex_color):
    """Applique une couleur de fond a une cellule de tableau."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_columns(document, num_cols):
    """Configure le nombre de colonnes de la section courante (comme Disposition > Colonnes)."""
    section = document.sections[-1]
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), '425')


def add_header_footer(document, header_text, add_page_numbering=True, different_first_page=False):
    section = document.sections[0]
    section.different_first_page_header_footer = different_first_page

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = header_text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if add_page_numbering:
        fp.add_run("Page ")
        add_field(fp, "PAGE", cached_text="2")
        fp.add_run(" sur ")
        add_field(fp, "NUMPAGES", cached_text="2")

    if different_first_page:
        # En-tete/pied de premiere page vides
        first_header = section.first_page_header
        first_header.paragraphs[0].text = ""
        first_footer = section.first_page_footer
        first_footer.paragraphs[0].text = ""


def title_page(doc, title, subtitle=""):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(16)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Exercice 1 - Mise en forme de caracteres et paragraphes
# ---------------------------------------------------------------------------

def ex1():
    doc = Document()

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Les Amis du Quartier")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    s = doc.add_paragraph()
    r = s.add_run("Notre mission")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.paragraph_format.line_spacing = 1.5
    p1.add_run(
        "Depuis 2020, notre association rassemble les habitants du quartier autour de projets "
        "solidaires. Nous croyons que "
    )
    r = p1.add_run("l'engagement bénévole")
    r.bold = True
    r.underline = True
    p1.add_run(
        " reste le meilleur moteur de changement local. "
    )
    r2 = p1.add_run("Chaque contribution compte, quelle que soit sa taille.")
    r2.font.highlight_color = 7  # jaune (WD_COLOR_INDEX.YELLOW)

    v = doc.add_paragraph()
    r = v.add_run("Nos trois valeurs :")
    r.bold = True

    for val in ["Solidarité", "Transparence", "Engagement durable"]:
        doc.add_paragraph(val, style="List Bullet")

    doc.save(os.path.join(OUT, "exercice1_solution.docx"))


# ---------------------------------------------------------------------------
# Exercice 2 - Liste multiniveaux
# ---------------------------------------------------------------------------

def ex2():
    doc = Document()
    doc.add_paragraph("Plan d'organisation — Fête de quartier annuelle", style="Heading 1")

    plan = [
        (0, "Avant l'événement"),
        (1, "Réserver la salle municipale"),
        (1, "Envoyer les invitations aux habitants"),
        (1, "Commander le matériel (chaises, sonorisation)"),
        (0, "Pendant l'événement"),
        (1, "Accueil et inscription des invités"),
        (1, "Animation et jeux pour les enfants"),
        (1, "Service de la collation"),
        (0, "Après l'événement"),
        (1, "Envoyer les remerciements"),
        (1, "Rédiger le bilan financier"),
    ]

    level1_num = 0
    level2_letters = "abcdefgh"
    level2_idx = 0
    for level, text in plan:
        if level == 0:
            level1_num += 1
            level2_idx = 0
            p = doc.add_paragraph()
            r = p.add_run(f"{level1_num}. {text}")
            r.bold = True
            r.font.size = Pt(13)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.25)
            letter = level2_letters[level2_idx]
            p.add_run(f"{letter}. {text}")
            level2_idx += 1

    doc.save(os.path.join(OUT, "exercice2_solution.docx"))


# ---------------------------------------------------------------------------
# Exercice 3 - Tabulations
# ---------------------------------------------------------------------------

def ex3():
    from docx.enum.text import WD_TAB_ALIGNMENT

    doc = Document()
    doc.add_paragraph("Fiche de contact bénévole", style="Heading 1")

    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run("Nom : Jean Baptiste\tTéléphone : 3712-3456")

    doc.add_paragraph()
    doc.add_paragraph("Liste de fournitures (taquet décimal)", style="Heading 2")

    items = [("Cahier", "25"), ("Stylo", "5,50"), ("Classeur", "120"), ("Ramette de papier", "310,75")]
    for name, price in items:
        p = doc.add_paragraph()
        ts = p.paragraph_format.tab_stops
        ts.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.DECIMAL)
        p.add_run(f"{name}\t{price}")

    doc.save(os.path.join(OUT, "exercice3_solution.docx"))


# ---------------------------------------------------------------------------
# Exercice 4 - Mise en page complete
# ---------------------------------------------------------------------------

def ex4():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    title_page(doc, "Rapport trimestriel", "Association Les Amis du Quartier")

    add_header_footer(doc, "Association Les Amis du Quartier", add_page_numbering=True, different_first_page=True)

    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph(
        "Ce rapport présente les activités du trimestre. Ce document illustre une mise en page "
        "avec en-tête, pied de page numéroté, et une page de garde sans en-tête ni numérotation."
    )
    doc.add_page_break()
    doc.add_paragraph("Activités du trimestre", style="Heading 1")
    doc.add_paragraph(
        "Trois événements majeurs ont marqué ce trimestre : la fête de quartier, la collecte "
        "alimentaire, et l'atelier de sensibilisation au recyclage."
    )

    doc.save(os.path.join(OUT, "exercice4_solution.docx"))


# ---------------------------------------------------------------------------
# Exercice 5 - Tableau avance
# ---------------------------------------------------------------------------

def ex5():
    doc = Document()
    doc.add_paragraph("Détail des dépenses trimestrielles", style="Heading 1")

    data = [("Fournitures", 450), ("Transport", 280), ("Communication", 120), ("Location de salle", 610)]
    data_sorted = sorted(data, key=lambda x: x[1], reverse=True)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Ligne de titre fusionnee
    hdr_cells = table.rows[0].cells
    merged = hdr_cells[0].merge(hdr_cells[1])
    merged.text = "Dépenses T1 2026"
    merged.paragraphs[0].runs[0].bold = True
    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(merged, "1B4F72")
    merged.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # En-tete de colonnes
    row = table.add_row().cells
    row[0].text = "Catégorie"
    row[1].text = "Montant (HTG)"
    for c in row:
        c.paragraphs[0].runs[0].bold = True
        set_cell_shading(c, "D6E4F0")

    total = 0
    for cat, montant in data_sorted:
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = str(montant)
        total += montant

    # Ligne de total avec vrai champ Word =SUM(ABOVE), valeur mise en cache
    # pour un rendu correct immediat (Word recalculera a l'ouverture / F9).
    row = table.add_row().cells
    p0 = row[0].paragraphs[0]
    p0.add_run("Total").bold = True
    p1 = row[1].paragraphs[0]
    add_field(p1, r'=SUM(ABOVE) \# "#,##0"', bold=True, cached_text=f"{total:,}".replace(",", " "))
    set_cell_shading(row[0], "F4D03F")
    set_cell_shading(row[1], "F4D03F")

    doc.add_paragraph()
    doc.add_paragraph(
        "(Ce total est un vrai champ Word =SUM(ABOVE) : si tu modifies un montant, fais un clic "
        "droit dessus puis « Mettre à jour les champs » pour le recalculer automatiquement.)"
    ).italic = True

    doc.save(os.path.join(OUT, "exercice5_solution.docx"))


# ---------------------------------------------------------------------------
# Exercice 6 - Styles personnalises + Table des matieres
# ---------------------------------------------------------------------------

def ex6():
    doc = Document()
    doc.add_paragraph("Rapport annuel — Sommaire", style="Heading 1")

    p = doc.add_paragraph()
    p.add_run("Table des matières générée automatiquement (mise à jour requise à l'ouverture, F9) :").italic = True

    toc_p = doc.add_paragraph()
    add_field(toc_p, r'TOC \o "1-3" \h \z \u',
               cached_text="[Table des matières — clic droit puis « Mettre à jour les champs »]")

    # Apercu illustratif du resultat attendu une fois le champ actualise dans Word
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    preview_title = doc.add_paragraph()
    preview_title.add_run("Aperçu du résultat attendu après actualisation :").italic = True
    preview_title.paragraph_format.space_before = Pt(12)

    toc_entries = [
        ("Résumé exécutif", "2", 0),
        ("Contexte général", "2", 1),
        ("Activités de l'année", "2", 0),
        ("Distribution de matériel scolaire", "2", 1),
        ("Formation des bénévoles", "3", 1),
        ("Perspectives", "3", 0),
    ]
    for text, page, level in toc_entries:
        tp = doc.add_paragraph()
        tp.paragraph_format.left_indent = Cm(0.5 * level)
        ts = tp.paragraph_format.tab_stops
        ts.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = tp.add_run(f"{text}\t{page}")
        if level == 0:
            run.bold = True

    doc.add_page_break()
    doc.add_paragraph("Résumé exécutif", style="Heading 1")
    doc.add_paragraph("Ce chapitre présente une synthèse des activités de l'année écoulée.")

    doc.add_paragraph("Contexte général", style="Heading 2")
    doc.add_paragraph("L'association a poursuivi sa mission de solidarité locale tout au long de l'année.")

    doc.add_paragraph("Activités de l'année", style="Heading 1")
    doc.add_paragraph("Trois grands axes ont structuré nos actions cette année.")

    doc.add_paragraph("Distribution de matériel scolaire", style="Heading 2")
    doc.add_paragraph("Plus de deux cents enfants ont bénéficié d'un kit scolaire complet.")

    doc.add_paragraph("Formation des bénévoles", style="Heading 2")
    doc.add_paragraph("Douze bénévoles ont suivi une formation à la gestion de projet associatif.")

    doc.add_paragraph("Perspectives", style="Heading 1")
    doc.add_paragraph("L'année prochaine verra le lancement d'un nouveau programme culturel.")

    doc.save(os.path.join(OUT, "exercice6_solution.docx"))


# ---------------------------------------------------------------------------
# Exercice 7 - Image avec legende
# ---------------------------------------------------------------------------

def ex7():
    doc = Document()
    doc.add_paragraph("Notre identité visuelle", style="Heading 1")
    doc.add_paragraph(
        "Le logo ci-dessous représente l'identité visuelle de notre association, utilisée sur "
        "tous nos documents officiels et notre correspondance."
    )

    logo_path = os.path.join(ASSETS, "logo_association.png")
    doc.add_picture(logo_path, width=Inches(3))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Figure 1 : Logo de l'association Les Amis du Quartier")
    r.italic = True
    r.font.size = Pt(10)

    doc.add_paragraph(
        "\nNote pédagogique : en classe, appliquez l'habillage « Rapproché » (onglet Format de "
        "l'image) pour que le texte environnant épouse le contour du logo, puis positionnez-le "
        "à droite du premier paragraphe — une manipulation interactive à réaliser directement "
        "dans Word (chapitre 21 du manuel)."
    ).italic = True

    doc.save(os.path.join(OUT, "exercice7_solution.docx"))


# ---------------------------------------------------------------------------
# Exercice 8 - Notes de bas de page (simulees) + bibliographie
# ---------------------------------------------------------------------------

def ex8():
    doc = Document()
    doc.add_paragraph("L'engagement bénévole en Haïti", style="Heading 1")

    p = doc.add_paragraph()
    p.add_run(
        "L'engagement bénévole reste le premier facteur de réussite d'un programme associatif "
        "local"
    )
    sup = p.add_run("1")
    sup.font.superscript = True
    p.add_run(
        ", en particulier dans les zones rurales où les ressources institutionnelles restent "
        "limitées"
    )
    sup2 = p.add_run("2")
    sup2.font.superscript = True
    p.add_run(".")

    doc.add_paragraph()
    note_title = doc.add_paragraph()
    note_title.add_run("Notes").bold = True
    n1 = doc.add_paragraph()
    n1.add_run("1. ").font.superscript = False
    n1.add_run(
        "Étude de référence sur l'engagement associatif, Institut Exemple, 2025, p. 12."
    ).font.size = Pt(9)
    n2 = doc.add_paragraph()
    n2.add_run("2. ")
    n2.add_run(
        "Observation de terrain menée par l'ONG Exemple International, mars 2026."
    ).font.size = Pt(9)

    doc.add_paragraph(
        "\n(Note pédagogique : en classe, utilisez de vraies notes de bas de page — Références > "
        "Insérer une note de bas de page, `Ctrl+Alt+F` — plutôt que cette simulation par exposant, "
        "qui illustre uniquement le rendu visuel attendu, chapitre 29 du manuel.)"
    ).italic = True

    doc.add_paragraph()
    doc.add_paragraph("Bibliographie", style="Heading 1")

    refs = [
        "Institut Exemple. (2025). Étude de référence sur l'engagement associatif. Port-au-Prince : Éditions Exemple.",
        "ONG Exemple International. (2026). Rapport annuel d'activités. Pignon : ONG Exemple International.",
    ]
    for ref in refs:
        rp = doc.add_paragraph()
        rp.paragraph_format.left_indent = Cm(1.25)
        rp.paragraph_format.first_line_indent = Cm(-1.25)
        rp.add_run(ref)

    doc.save(os.path.join(OUT, "exercice8_solution.docx"))


# ---------------------------------------------------------------------------
# Exercice 9 - Colonnes + lettrine (approximee)
# ---------------------------------------------------------------------------

def ex9():
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Bulletin associatif — Édition de printemps")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    doc.add_paragraph()

    # Nouvelle "section" en 2 colonnes pour le corps de l'article
    doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(doc, 2)

    p = doc.add_paragraph()
    # Lettrine approximee : premiere lettre agrandie
    first_letter = p.add_run("N")
    first_letter.font.size = Pt(40)
    first_letter.bold = True
    first_letter.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    rest = p.add_run(
        "otre association a le plaisir de vous présenter, dans ce numéro de printemps, un "
        "bilan complet de nos activités récentes ainsi que nos projets pour la saison à venir. "
        "Grâce à l'engagement constant de nos bénévoles, plusieurs initiatives ont vu le jour "
        "cette année, touchant directement plus de trois cents familles du quartier."
    )
    p.paragraph_format.line_spacing = 1.0

    doc.add_paragraph(
        "Parmi les temps forts, la fête de quartier annuelle a rassemblé un public record, "
        "confirmant l'attachement de la communauté à ces moments de partage collectif. "
        "L'atelier de sensibilisation au recyclage a quant à lui permis de former une vingtaine "
        "de jeunes bénévoles à des gestes simples mais efficaces."
    )

    # Retour a une seule colonne pour la suite
    doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(doc, 1)
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "(Note pédagogique : la grande lettre initiale ci-dessus est une approximation visuelle. "
        "En classe, utilisez la vraie fonctionnalité Lettrine — Insertion > Lettrine, chapitre 16 "
        "du manuel — qui intègre la lettre dans le flux du texte environnant.)"
    ).italic = True

    doc.save(os.path.join(OUT, "exercice9_solution.docx"))


# ---------------------------------------------------------------------------
# Exercice 10 - Publipostage : exemple de lettre fusionnee
# ---------------------------------------------------------------------------

def ex10():
    doc = Document()
    doc.add_paragraph("Pignon, le 4 août 2026\n").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("Madame Marie Delva\n123 Rue de l'Exemple\nPignon, Haïti\n")

    doc.add_paragraph("Chère Madame Delva,")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "Au nom de toute l'équipe de l'ONG Exemple International, nous tenons à vous exprimer "
        "notre plus sincère reconnaissance pour votre don de "
    )
    r = p.add_run("5 000 HTG")
    r.bold = True
    p.add_run(
        " reçu le mois dernier. Votre générosité exceptionnelle a un impact direct sur nos "
        "programmes de distribution de matériel scolaire, qui bénéficient chaque mois à des "
        "dizaines de familles de notre communauté."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Grâce à des donateurs fidèles comme vous, nous pouvons continuer à porter nos actions "
        "solidaires sur le terrain, avec constance et détermination."
    )
    doc.add_paragraph()
    doc.add_paragraph("Merci à toute notre équipe bénévole pour son engagement continu à vos côtés.")
    doc.add_paragraph()
    doc.add_paragraph("La Coordonnatrice,\nMarie-Ange Joseph")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "(Note pédagogique : cette lettre représente le résultat de la fusion pour UN seul "
        "destinataire de la liste. En classe, chaque étudiant doit produire la fusion complète "
        "pour les trois destinataires fournis — chapitre 34 du manuel — avec un paragraphe "
        "conditionnel différent selon que le don dépasse ou non 1000 HTG.)"
    ).italic = True

    doc.save(os.path.join(OUT, "exercice10_solution.docx"))


if __name__ == "__main__":
    fns = [ex1, ex2, ex3, ex4, ex5, ex6, ex7, ex8, ex9, ex10]
    for fn in fns:
        fn()
        print(f"{fn.__name__} : OK")
    print("Tous les fichiers solutions .docx ont ete generes dans", OUT)
